#!/usr/bin/env python3
"""Convert alt-text manifest Markdown into a standard-format .docx for layout hand-off.

Usage:
    python manifest_to_docx.py <manifest.md> [--out <path.docx>] [--subtitle "<text>"]
    python manifest_to_docx.py <dir> [--subtitle "<text>"]      # recurse: every *alt-text-manifest.md

    # optionally embed a thumbnail beside each description (needs Pillow):
    python manifest_to_docx.py <manifest.md> --out <path.docx> --thumbs \
        --images-dir "<folder with the images>" [--raster-dir "<rasterized PDF pages>"] \
        [--thumb-width 1.6]

Per image the .docx shows: Category, Alt text, Long description (complex only), and a Note
ONLY when one is present. Title = project name from the manifest heading; subtitle = --subtitle.

Requires python-docx. --thumbs additionally requires Pillow. See README.md for the venv bootstrap.
"""
import os, re, sys, glob, argparse
from io import BytesIO
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.table import WD_ALIGN_VERTICAL
except ImportError:
    sys.exit("python-docx is required: python3 -m venv .venv && .venv/bin/pip install python-docx")
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from manifest_common import (parse_manifest, clean_note, norm_cat, norm_key,
                             project_name, resolve_image_path)

DEFAULT_SUBTITLE = "Image Alt Text"

# ---- docx helpers ----
def add_hr(paragraph, color="888888", sz="8"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr'); bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), sz)
    bottom.set(qn('w:space'), '2'); bottom.set(qn('w:color'), color)
    pBdr.append(bottom); pPr.append(pBdr)

def _fill_field(p, label, value, color=None):
    r = p.add_run(label); r.bold = True
    if color:
        r.font.color.rgb = color
    p.add_run(" "); p.add_run(value)
    p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(0)

def set_cell_width(cell, inches):
    cell.width = Inches(inches)

def make_thumb(path, max_px):
    """Downscale (and EXIF-rotate) an image to a PNG stream; None on failure."""
    try:
        from PIL import Image, ImageOps
    except ImportError:
        sys.exit("--thumbs requires Pillow: .venv/bin/pip install Pillow")
    try:
        im = Image.open(path)
        im = ImageOps.exif_transpose(im)          # honor camera/source rotation
        if im.mode not in ("RGB", "L"):
            im = im.convert("RGB")
        im.thumbnail((max_px, max_px))
        bio = BytesIO(); im.save(bio, format="PNG"); bio.seek(0)
        return bio
    except Exception:
        return None

def write_entry_text(container, row, model, first_para=None):
    """Write Category / Alt text / Long description / Note into a doc or table cell."""
    cat = norm_cat(row['category']); alt = row['alt'].strip()
    note = clean_note(row['notes_raw'])
    long_paras = model['longs'].get(norm_key(row['file']))

    p = first_para if first_para is not None else container.add_paragraph()
    _fill_field(p, "Category:", cat)
    _fill_field(container.add_paragraph(), "Alt text:",
                alt if alt else ('""' if cat == 'Decorative' else alt))
    if long_paras:
        lp = container.add_paragraph(); lp.paragraph_format.space_after = Pt(2)
        lp.add_run("Long description:").bold = True
        for para in long_paras:
            dp = container.add_paragraph(para)
            dp.paragraph_format.left_indent = Inches(0.2); dp.paragraph_format.space_after = Pt(4)
    if note:
        _fill_field(container.add_paragraph(), "Note:", note, RGBColor(0x99, 0x5A, 0x00))
    return long_paras is not None

def build_doc(model, out_path, subtitle, thumbs=False, images_dir=None,
              raster_dir=None, thumb_width=1.6):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Inches(1))
    normal = doc.styles['Normal']; normal.font.name = 'Calibri'; normal.font.size = Pt(11)
    style_names = [s.name for s in doc.styles]
    grey = RGBColor(0x60, 0x60, 0x60)
    content_w = 6.5  # letter minus 1" margins
    left_w = min(max(thumb_width + 0.2, 1.0), 3.0)
    right_w = content_w - left_w

    doc.add_heading(project_name(model['title']), level=0)
    if subtitle:
        sp = doc.add_paragraph(subtitle)
        if 'Subtitle' in style_names:
            sp.style = doc.styles['Subtitle']
    def meta(text):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text); r.italic = True; r.font.size = Pt(9); r.font.color.rgb = grey
        return p
    if model['people']:
        meta(f"People policy: {model['people']}")
    if model['generated']:
        meta(f"Generated: {model['generated']}")
    legend = meta("For each image below: Category, Alt text, Long description (complex figures only), "
                  "and a Note where one applies. Alt text is the caption-ready text for InDesign.")
    add_hr(legend); legend.paragraph_format.space_after = Pt(8)

    used = set(); missing_thumbs = []
    for row in model['rows']:
        head = doc.add_heading(f"Image {row['idx']}.  {row['file']}", level=2)
        head.paragraph_format.space_before = Pt(10); head.paragraph_format.space_after = Pt(3)

        img_path = None
        if thumbs:
            img_path = resolve_image_path(row['file'], images_dir, raster_dir)
        stream = make_thumb(img_path, int(thumb_width * 220)) if img_path else None

        if stream is not None:
            table = doc.add_table(rows=1, cols=2)
            table.autofit = False; table.allow_autofit = False
            left, right = table.rows[0].cells
            set_cell_width(left, left_w); set_cell_width(right, right_w)
            left.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            right.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            left.paragraphs[0].add_run().add_picture(stream, width=Inches(thumb_width))
            write_entry_text(right, row, model, first_para=right.paragraphs[0])
        else:
            if thumbs and img_path is None:
                missing_thumbs.append(row['file'])
            write_entry_text(doc, row, model)

        if model['longs'].get(norm_key(row['file'])) is not None:
            used.add(norm_key(row['file']))

    orphan = [(k, v) for k, v in model['longs'].items() if k not in used]
    if orphan:
        doc.add_heading("Additional long descriptions", level=2)
        for k, paras in orphan:
            doc.add_paragraph(k).runs[0].bold = True
            for para in paras:
                doc.add_paragraph(para)

    doc.save(out_path)
    return len(model['rows']), len(orphan), missing_thumbs

def iter_manifests(paths):
    for p in paths:
        if os.path.isdir(p):
            for m in sorted(set(glob.glob(os.path.join(p, "**", "*alt-text-manifest.md"), recursive=True)
                                + glob.glob(os.path.join(p, "*alt-text-manifest.md")))):
                yield m
        else:
            yield p

def main():
    ap = argparse.ArgumentParser(description="Convert alt-text manifest .md to a standard .docx.")
    ap.add_argument("paths", nargs="+", help="manifest .md file(s) or dir(s) to recurse")
    ap.add_argument("--out", help="output .docx path (only valid with a single manifest file)")
    ap.add_argument("--subtitle", default=DEFAULT_SUBTITLE, help="report label shown under the title")
    ap.add_argument("--thumbs", action="store_true", help="embed a thumbnail beside each entry (needs Pillow)")
    ap.add_argument("--images-dir", help="folder holding the images (default: the manifest's own folder)")
    ap.add_argument("--raster-dir", help="folder holding rasterized PDF pages (default: --images-dir)")
    ap.add_argument("--thumb-width", type=float, default=1.6, help="thumbnail width in inches (default 1.6)")
    args = ap.parse_args()

    items = list(iter_manifests(args.paths))
    if args.out and len(items) != 1:
        sys.exit("--out is only valid when exactly one manifest is given")
    total = 0
    for md in items:
        model = parse_manifest(md)
        out = args.out if args.out else (md[:-3] + ".docx")
        images_dir = args.images_dir or os.path.dirname(os.path.abspath(md))
        rows, orphan, missing = build_doc(model, out, args.subtitle, thumbs=args.thumbs,
                                          images_dir=images_dir, raster_dir=args.raster_dir,
                                          thumb_width=args.thumb_width)
        flags = ""
        if orphan: flags += f"  [!! {orphan} orphan long-desc]"
        if missing: flags += f"  [{len(missing)} thumb(s) not found]"
        print(f"OK  rows={rows:2d}{flags}  ->  {out}")
        total += 1
    print(f"\nGenerated {total} .docx file(s).")

if __name__ == "__main__":
    main()
